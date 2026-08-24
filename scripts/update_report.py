from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK
from docx.shared import Pt


SOURCE = Path(r"E:\APPS\báo cáo bài tập lớn\báo cáo android.docx")
OUTPUT = Path(r"E:\taskmanager\BaoCao_Android_CapNhat.docx")


def set_cell(table, row, column, value):
    table.rows[row].cells[column].text = value


def remove_table(table):
    element = table._element
    element.getparent().remove(element)


def remove_row(table, row):
    table._tbl.remove(row._tr)


document = Document(SOURCE)
paragraphs = document.paragraphs

paragraph_updates = {
    7: "Hỗ trợ phân cấp trách nhiệm Owner → Admin → Member; Owner chịu trách nhiệm cao nhất, Admin hỗ trợ quản lý và Member thực hiện công việc.",
    8: "Cung cấp Dashboard số liệu tổng hợp và biểu đồ Gantt có trục ngày, thanh thời gian bắt đầu–hạn chót và lớp tiến độ.",
    12: "Ứng dụng Android native (Java) theo hướng offline-first: SQLite lưu cục bộ, Firebase Authentication quản lý tài khoản và Cloud Firestore đồng bộ nhiều thiết bị.",
    14: "Giao việc cho một hoặc nhiều thành viên, checklist tự tính tiến độ, Dashboard và biểu đồ Gantt.",
    16: "Đồng bộ dữ liệu gần thời gian thực giữa nhiều thiết bị bằng Cloud Firestore, có hàng đợi đồng bộ khi thiết bị ngoại tuyến.",
    17: "Bình luận trong công việc đã triển khai; tệp đính kèm cloud được tắt trong chế độ Firebase Spark và là hướng nâng cấp sau.",
    18: "Thông báo lịch cục bộ hoạt động khi đến giờ bắt đầu, trước hạn một giờ và khi quá hạn; push notification từ thiết bị khác là hướng phát triển cần Cloud Functions.",
    19: "Đăng nhập mạng xã hội và OTP tùy biến nằm ngoài phạm vi hiện tại; quên mật khẩu dùng liên kết bảo mật do Firebase gửi qua email.",
    25: "Xem và cập nhật thông tin cá nhân; xem mã công khai duy nhất dạng USR-… để mời vào Team mà không công khai Firebase UID.",
    29: "Tiến độ được tính tự động từ số bước nhỏ đã hoàn thành; hệ thống cảnh báo lúc bắt đầu, sắp hết hạn và quá hạn.",
    31: "Tạo nhóm; người tạo tự động trở thành Owner của nhóm.",
    32: "Mời thành viên tham gia nhóm bằng mã công khai USR-…; người được mời có thể chấp nhận hoặc từ chối trên thiết bị khác.",
    33: "Owner có thể cấp quyền Admin hoặc đưa người dùng về vai trò Member; Admin quản lý thành viên trong phạm vi được cho phép.",
    36: "Owner/Admin/Member có thể tạo công việc theo quyền nghiệp vụ và giao cho một hoặc nhiều thành viên đang hoạt động trong nhóm.",
    37: "Mọi thành viên được giao có thể cập nhật trạng thái và checklist; thanh tiến độ tự tăng theo số bước đã hoàn thành.",
    41: "Xem biểu đồ Gantt thật theo trục ngày, thời gian bắt đầu–hạn chót và phần trăm hoàn thành của từng công việc.",
    44: "Bảo mật: mật khẩu chỉ do Firebase Authentication quản lý, không lưu trong SQLite/Firestore; quyền truy cập được kiểm tra ở Repository và Firestore Security Rules.",
    47: "Khả năng bảo trì và mở rộng: kiến trúc tách giao diện, ViewModel, Repository, DAO/SQLite và lớp đồng bộ Firebase để có thể nâng cấp giao diện mà không đổi nghiệp vụ.",
    49: "Ràng buộc triển khai: SQLite hoạt động offline-first và Cloud Firestore đồng bộ nhiều tài khoản/thiết bị thật. Bản hiện tại dùng gói Spark miễn phí, không phụ thuộc Storage hoặc Cloud Functions.",
    51: "Hệ thống gồm bốn nhóm Actor: User cá nhân, Owner, Admin và Member trong Team.",
    97: "Biểu đồ trình tự tạo dự án và chỉ định người quản lý",
    105: "Biểu đồ trình tự AI hỗ trợ (định hướng tương lai, chưa triển khai trong sản phẩm hiện tại)",
    133: "Workspace PERSONAL thuộc người tạo; trường manager_id chỉ dùng cho Workspace TEAM.",
    134: "Workspace TEAM phải có manager_id trỏ tới tài khoản giữ vai trò Owner.",
    136: "Khi chuyển quyền Owner phải cập nhật manager_id và vai trò trong WORKSPACE_MEMBERS bằng transaction.",
    139: "OWNER",
    140: "ADMIN",
    141: "MEMBER",
    142: "Giá trị status:",
    143: "ACTIVE",
    144: "REMOVED",
    145: "Owner/Admin/Member được lưu trực tiếp trong WORKSPACE_MEMBERS; Owner là duy nhất trong mỗi Team.",
    146: "Bảng TEAM_INVITES",
    150: "REJECTED",
    154: "invited_by phải là Owner hoặc Admin có quyền quản lý thành viên.",
    159: "4.5. Bảng PROJECTS",
    160: "Giá trị status:",
    161: "ACTIVE",
    162: "COMPLETED",
    163: "ARCHIVED",
    164: "Quy tắc:",
    165: "Dự án chỉ thuộc workspace loại TEAM.",
    166: "created_by phải là thành viên ACTIVE có quyền tạo dự án.",
    167: "manager_id, nếu có, phải là thành viên ACTIVE của Team.",
    168: "due_date phải lớn hơn start_date.",
    169: "Tiến độ dự án được tổng hợp từ TASKS và checklist, không nhập thủ công độc lập.",
    170: "Dự án có thể chứa mốc thời gian (milestone) và nhiều công việc.",
    171: "Không xóa cứng dự án đang có dữ liệu; dùng trạng thái ARCHIVED.",
    172: "Mọi thay đổi được đồng bộ theo version và sync_status.",
    173: "",
    186: "progress nằm trong khoảng 0–100 và được tính từ TASK_SUBTASKS khi công việc có checklist.",
    192: "4.7. Bảng TASK_ASSIGNEES",
    194: "Quy tắc:",
    195: "Người được giao phải là Member ACTIVE trong cùng Team.",
    196: "Một Task có thể được giao cho nhiều User; khóa chính ghép task_id + user_id ngăn phân công trùng.",
    197: "Công việc cá nhân không cần bản ghi phân công.",
    198: "assigned_by là người thực hiện thao tác phân công.",
    199: "Khi sửa phân công, ứng dụng cập nhật transaction để tránh trạng thái dở dang.",
    200: "Dữ liệu cloud duy trì assigneeIds và assigneeId đầu tiên để tương thích bản cũ.",
    201: "",
    202: "",
    217: "Tác giả được sửa/xóa mềm bình luận của mình; Owner/Admin có thể xóa mềm bình luận theo quyền quản lý.",
    237: "4.11. Bảng NOTIFICATIONS",
    254: "4.12. Bảng AI_INTERACTIONS (định hướng tương lai, chưa triển khai)",
    263: "4.13. Bảng USER_DEVICES (chỉ dùng khi triển khai FCM/Cloud Functions)",
    265: "4.14. Bảng SYNC_QUEUE",
    280: "DEPENDENCY",
    283: "Chương 3: Kết quả sản phẩm",
}

for index, value in paragraph_updates.items():
    paragraphs[index].text = value

# Bỏ toàn bộ phần mô hình bấm giờ vì sản phẩm không còn chức năng này.
for index in range(224, 217, -1):
    element = paragraphs[index]._element
    element.getparent().remove(element)

# Cập nhật mô tả Actor/use case và các trường dữ liệu quan trọng.
tables = document.tables
table_updates = {
    (0, 2, 0): "Owner",
    (0, 2, 1): "Sở hữu Team, quản lý thành viên, dự án và theo dõi tiến độ tổng thể",
    (0, 3, 0): "Admin",
    (0, 3, 1): "Hỗ trợ quản lý thành viên, dự án và công việc theo quyền",
    (4, 4, 1): "Nhóm mới được tạo (có Owner); hoặc người dùng chính thức tham gia/không tham gia Team được mời",
    (5, 1, 1): "Owner/Admin",
    (5, 2, 1): "Mời bằng mã USR-…, đổi vai trò, xóa thành viên hoặc giải tán/rời Team theo quyền",
    (5, 3, 1): "Đã đăng nhập và có quyền quản lý thành viên",
    (6, 1, 1): "Owner/Admin",
    (6, 4, 1): "Dự án được tạo/cập nhật, có người quản lý và mốc thời gian rõ ràng",
    (7, 1, 1): "Owner/Admin/Member đang hoạt động",
    (7, 2, 1): "Tạo công việc trong dự án và giao cho một hoặc nhiều thành viên",
    (9, 1, 1): "Mọi thành viên Team; Owner/Admin có số liệu quản trị toàn nhóm",
    (11, 2, 1): "Owner",
    (11, 2, 2): "Sở hữu Team, quản lý thành viên và dự án",
    (11, 3, 1): "Admin",
    (11, 3, 2): "Hỗ trợ quản lý dự án, thành viên và phân công",
    (12, 2, 0): "Owner",
    (12, 2, 1): "Thông tin nhóm, lời mời bằng mã USR-…, dự án, phân công và yêu cầu báo cáo",
    (12, 3, 0): "Admin",
    (12, 3, 1): "Dữ liệu dự án, thành viên, phân công và yêu cầu thống kê",
    (12, 4, 1): "Trạng thái, checklist, bình luận và kết quả công việc",
    (13, 4, 1): "Quản lý dự án và mốc thời gian",
    (13, 4, 2): "Tạo/cập nhật/lưu trữ dự án, đặt lịch và quản lý milestone",
    (13, 5, 2): "Tạo, sửa, xóa và giao công việc cho một hoặc nhiều thành viên",
    (13, 6, 2): "Cập nhật checklist; tiến độ tự tính và phát sinh nhắc việc theo lịch",
    (13, 7, 1): "Quản lý thông báo",
    (13, 7, 2): "Lưu thông báo trong app và lập lịch nhắc cục bộ",
    (14, 4, 2): "Thông tin dự án, người quản lý và milestone",
    (14, 6, 2): "Checklist, bình luận, phụ thuộc và lịch sử thay đổi",
    (14, 7, 1): "Thông báo",
    (14, 7, 2): "Thông báo trong app và dữ liệu lịch nhắc",
    (16, 3, 3): "Owner của workspace Team",
    (18, 3, 3): "Owner/Admin gửi lời mời",
    (19, 3, 3): "Thành viên tạo dự án",
    (19, 4, 0): "manager_id",
    (19, 4, 3): "Thành viên quản lý dự án (tùy chọn)",
    (19, 9, 0): "due_date",
    (21, 2, 3): "Thành viên được giao; một Task có thể có nhiều dòng",
    (21, 3, 3): "Người thực hiện phân công",
}

for (table_index, row, column), value in table_updates.items():
    set_cell(tables[table_index], row, column, value)

# Bỏ bảng WORK_SESSIONS và quan hệ tương ứng.
for table in list(document.tables):
    if any("session_id" in cell.text for row in table.rows for cell in row.cells):
        remove_table(table)

for table in document.tables:
    for row in list(table.rows):
        if any("WORK_SESSIONS" in cell.text for cell in row.cells):
            remove_row(table, row)

# Đánh dấu rõ bảng AI chỉ là thiết kế tương lai.
for table in document.tables:
    if any("interaction_id" in cell.text for row in table.rows for cell in row.cells):
        table.rows[0].cells[3].text = "Ý nghĩa (định hướng tương lai)"

# Viết Chương 3 theo đúng các đối tượng và chức năng đã triển khai.
document.add_heading("3.1. Tài khoản và xác thực", level=2)
document.add_paragraph(
    "Người dùng có các màn hình trang chủ, đăng ký, đăng nhập, quên mật khẩu và hồ sơ dùng chung. "
    "Firebase Authentication xử lý đăng ký/đăng nhập, xác minh email và gửi liên kết đặt lại mật khẩu. "
    "Mỗi tài khoản có mã USR-… duy nhất để mời vào Team; ảnh đại diện ở chế độ Spark dùng chữ cái tên người dùng."
)

document.add_heading("3.2. Chức năng công việc cá nhân", level=2)
document.add_paragraph(
    "Người dùng tạo công việc với tiêu đề, mô tả, ưu tiên, ngày giờ bắt đầu và hạn hoàn thành cụ thể. "
    "Mỗi công việc có nhiều bước nhỏ; khi đánh dấu hoàn thành một bước, tiến độ và trạng thái công việc lớn được tính lại tự động. "
    "Danh sách có tìm kiếm, lọc, Dashboard cá nhân, lịch sử thay đổi, thùng rác và khôi phục."
)

document.add_heading("3.3. Owner của Team", level=2)
document.add_paragraph(
    "Owner tạo/chỉnh sửa/giải tán Team, mời thành viên bằng mã USR-…, đổi vai trò, xóa thành viên và chuyển quyền sở hữu. "
    "Owner quản lý toàn bộ dự án, milestone, công việc, phân công và Dashboard của Team. "
    "Các thao tác nhạy cảm đều được kiểm tra ở Repository và Firestore Security Rules."
)

document.add_heading("3.4. Admin và Member", level=2)
document.add_paragraph(
    "Admin hỗ trợ Owner quản lý thành viên, dự án, milestone và nội dung cộng tác theo quyền. "
    "Member xem dữ liệu Team, tạo công việc khi đang hoạt động và cập nhật các công việc được giao. "
    "Mọi vai trò đều có thể rời Team theo quy tắc; Owner phải chuyển quyền trước khi rời."
)

document.add_heading("3.5. Dự án và công việc nhiều người thực hiện", level=2)
document.add_paragraph(
    "Dự án có trạng thái, người quản lý, ngày giờ bắt đầu, hạn hoàn thành và các mốc milestone. "
    "Mỗi công việc có thể chọn đồng thời nhiều thành viên; tất cả người được giao đều được nhận lịch nhắc và có quyền cập nhật checklist. "
    "Ứng dụng giữ tương thích dữ liệu cũ bằng assigneeId đầu tiên đồng thời đồng bộ danh sách assigneeIds."
)

document.add_heading("3.6. Cộng tác và quan hệ phụ thuộc", level=2)
document.add_paragraph(
    "Thành viên có thể thêm, sửa và xóa mềm bình luận theo quyền tác giả/quản trị. "
    "Quan hệ phụ thuộc chỉ được tạo giữa các công việc cùng dự án, không được trùng hoặc tự phụ thuộc. "
    "Thuật toán duyệt toàn bộ đồ thị chặn cả vòng lặp gián tiếp; quan hệ có thể xóa và đồng bộ khi thiết bị trực tuyến trở lại."
)

document.add_heading("3.7. Dashboard và biểu đồ Gantt", level=2)
document.add_paragraph(
    "Dashboard hiển thị tổng số công việc, cần làm, đang làm, hoàn thành, đã hủy, quá hạn, tiến độ theo dự án và khối lượng theo thành viên. "
    "Biểu đồ Gantt dùng trục ngày thật, thanh từ ngày bắt đầu tới hạn chót và lớp màu thể hiện tiến độ. "
    "Danh sách công việc tải theo đợt 20 bản ghi và hỗ trợ bộ lọc dự án, thành viên, trạng thái."
)

document.add_heading("3.8. Thông báo và làm việc ngoại tuyến", level=2)
document.add_paragraph(
    "AlarmManager lập lịch thông báo khi đến giờ bắt đầu, trước hạn một giờ và sau khi quá hạn; thông báo cũng được lưu trong ứng dụng. "
    "SQLite cho phép tiếp tục làm việc khi mất mạng, còn SYNC_QUEUE đẩy các thay đổi lên Firestore khi kết nối trở lại. "
    "Firestore listeners cập nhật dữ liệu giữa các thiết bị và được giới hạn theo workspace."
)

document.add_heading("3.9. Phạm vi bản Spark và hướng nâng cấp", level=2)
document.add_paragraph(
    "Bản hiện tại không yêu cầu thanh toán: Authentication, Firestore, đồng bộ Team và nhắc giờ cục bộ hoạt động trên Firebase Spark. "
    "Storage/tệp đính kèm và Cloud Functions/push notification từ thiết bị khác đang tắt; đây là hai phần chỉ bật khi có ngân sách. "
    "Giao diện có thể tiếp tục nâng cấp Material 3, dark mode và responsive mà không phải thay đổi mô hình dữ liệu hoặc nghiệp vụ."
)

document.add_heading("3.10. Trạng thái kiểm thử", level=2)
status_table = document.add_table(rows=1, cols=3)
status_table.style = "Table Grid"
headers = status_table.rows[0].cells
headers[0].text = "Hạng mục"
headers[1].text = "Trạng thái"
headers[2].text = "Ghi chú"
rows = [
    ("Unit test Android", "Đạt", "Quy tắc tiến độ, lịch, vai trò, Gantt và vòng lặp phụ thuộc"),
    ("Firestore Rules Emulator", "Đạt", "Quyền nhiều tài khoản, bình luận, dependency và attachment metadata"),
    ("Build/Lint/Instrumentation", "Đạt", "Biên dịch APK và kiểm tra schema SQLite v7 trên Pixel emulator"),
    ("Hai thiết bị thật", "Cần nghiệm thu thủ công", "Thực hiện theo E2E_TEST_PLAN.md với hai email đã xác minh"),
]
for name, state, note in rows:
    cells = status_table.add_row().cells
    cells[0].text = name
    cells[1].text = state
    cells[2].text = note

document.add_paragraph()
note = document.add_paragraph()
run = note.add_run("Lưu ý: ")
run.bold = True
note.add_run(
    "Tài liệu này phản ánh bản triển khai ngày 23/08/2026. Các phần AI, Storage và Cloud Functions được ghi rõ là hướng phát triển, không tính là chức năng đã hoàn thành."
)

document.core_properties.title = "Báo cáo Android Task Manager – bản cập nhật"
document.core_properties.subject = "Ứng dụng quản lý công việc cá nhân và Team"
document.core_properties.modified = datetime.now(timezone.utc)

# Chuẩn hóa cỡ chữ cho nội dung Chương 3 mới, không thay đổi định dạng các phần cũ.
for paragraph in document.paragraphs[-35:]:
    for run in paragraph.runs:
        if run.font.size is None:
            run.font.size = Pt(11)

document.save(OUTPUT)
print(OUTPUT)
